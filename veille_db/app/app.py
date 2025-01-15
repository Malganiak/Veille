# app.py

import streamlit as st
from utils import *
import pandas as pd
import os
import requests
from bs4 import BeautifulSoup
from dotenv import load_dotenv
import pypdf
from docx import Document
from datetime import datetime

# Configuration de la page
st.set_page_config(layout="wide")

# CSS pour les bulles de dialogue
st.markdown("""
    <style>
    .user-message {
        background-color: #dcf8c6;
        border-radius: 10px;
        padding: 10px;
        margin: 10px 0;
        max-width: 80%;
        word-wrap: break-word;
        align-self: flex-end;
        width: 100%;
    }
    .bot-message {
        background-color: #f1f0f0;
        border-radius: 10px;
        padding: 10px;
        margin: 10px 0;
        max-width: 80%;
        word-wrap: break-word;
        align-self: flex-start;
        width: 100%;
    }
    .loading {
        font-size: 1.2em;
        color: #888;
        margin: 10px 0;
    }
    </style>
    """, unsafe_allow_html=True)

load_dotenv()

default_source_urls = load_default_sources()
default_keywords = load_default_keywords()
default_urls_summary = [
    "https://www.bearingpoint.com/fr-fr/publications-evenements/blogs/marketing-vente/exp%C3%A9rience-client-comment-ia-g%C3%A9n%C3%A9rative-va-renforcer-emotion/",
    "https://actualites.td.com/ca/fr/news/2024-05-16-la-td-lance-de-nouveaux-projets-pilotes-d-27intelligence-artif",
]
default_urls_corpus = [
    "https://www.cscience.ca/banque-cibc-un-projet-pilote-dia-generative-pour-le-service-a-la-clientele/",
    "https://reglo.ai/projet-ia-generative-points-de-vigilance/",
    "https://www.cio-online.com/actualites/lire-relation-client-best-western-confie-les-clefs-a-l-ia-generative-15745.html",
]

st.title("Automatisation de la veille")

tabs = st.tabs(
    [
        "Ma veille personnalisée",
        "Suggestions d'articles par thèmes personnalisés",
        "Suggestions d'articles par urls sources",
        "Résumé automatique d'article",
        "Synthèse de corpus d'articles",
        "Chatbot Q/A",
    ]
)

##################################
# Onglet 1 : Ma veille personnalisée
##################################
with tabs[0]:
    st.write("### Ma veille personnalisée")

    try:
        # Charger les sources et thèmes par défaut
        default_sources = load_default_sources()
        default_keywords = load_default_keywords()

        # Interface pour modifier les sources par défaut
        st.write("### Modifier les sources par défaut")
        new_sources = st.text_area("Entrez les nouvelles sources (une par ligne)", value="\n".join(default_sources))
        if st.button("Sauvegarder les sources"):
            try:
                save_default_sources(new_sources.split("\n"))
                st.success("Sources sauvegardées avec succès.")
            except Exception as e:
                st.error(f"Erreur lors de la sauvegarde des sources : {str(e)}")

        # Interface pour modifier les thèmes par défaut
        st.write("### Modifier les thèmes par défaut")
        new_keywords = st.text_area("Entrez les nouveaux thèmes (un par ligne)", value="\n".join(default_keywords))
        if st.button("Sauvegarder les thèmes"):
            try:
                save_default_keywords(new_keywords.split("\n"))
                st.success("Thèmes sauvegardés avec succès.")
            except Exception as e:
                st.error(f"Erreur lors de la sauvegarde des thèmes : {str(e)}")

    except Exception as e:
        st.error(f"Erreur de connexion à l'API : {str(e)}")

    # Interface pour gérer le temps de parution
    st.write("### Gérer le temps de parution")
    time_unit = st.selectbox("Unité de temps", ["mois", "années"])
    time_value = st.number_input("Paru il y a", min_value=1, step=1)

    # Interface pour gérer les filtres Google
    st.write("### Gérer les filtres Google")
    exclude_ads = st.checkbox("Exclure les pubs et contenus sponsorisés")
    exclude_professional = st.checkbox("Exclure les contenus professionnels")
    target_press = st.checkbox("Cibler uniquement les articles de presse")
    exclude_jobs = st.checkbox("Exclure les offres d'emploi")
    exclude_training = st.checkbox("Exclure les formations")

    # Sauvegarder les paramètres de filtres
    if st.button("Sauvegarder les filtres"):
        save_filters({
            "exclude_ads": exclude_ads,
            "exclude_professional": exclude_professional,
            "target_press": target_press,
            "time_unit": time_unit,
            "time_value": time_value,
            "exclude_jobs": exclude_jobs,
            "exclude_training": exclude_training,
        })
        st.success("Filtres sauvegardés avec succès.")

##################################
# Onglet 2 : Suggestions d'articles par thèmes personnalisés
##################################
with tabs[1]:
    st.write("### Suggestions d'articles par thèmes personnalisés")
    keywords = load_default_keywords()
    filters = load_filters()
    time_unit = filters.get("time_unit", "mois")
    time_value = filters.get("time_value", 1)
    input_data = "\n".join(keywords) + f"{time_unit}{time_value}"

    # Vérifier s'il y a des résultats précédemment sauvegardés
    if check_and_load_results(input_data, "summaries"):
        st.success("Chargement des résultats précédents.")
    else:
        if not keywords:
            st.error("Veuillez fournir au moins un mot-clé.")
        else:
            progress_bar = st.progress(0)
            status_text = st.empty()
            with st.spinner("Recherche des articles en cours..."):
                urls = []
                keyword_sources = []
                for idx, keyword in enumerate(keywords):
                    keyword_urls = google_search(
                        query=keyword,
                        num_results=st.session_state.get("num_articles_keywords", 10),
                        time_unit=time_unit,
                        time_value=time_value,
                        exclude_ads=filters.get("exclude_ads", False),
                        exclude_professional=filters.get("exclude_professional", False),
                        target_press=filters.get("target_press", False),
                        exclude_jobs=filters.get("exclude_jobs", False),
                        exclude_training=filters.get("exclude_training", False),
                    )
                    urls.extend(keyword_urls)
                    keyword_sources.extend([keyword] * len(keyword_urls))
                    progress_bar.progress((idx + 1) / len(keywords))
                    status_text.text(f"Recherche en cours... {idx + 1}/{len(keywords)}")

                if urls:
                    st.session_state["proposed_urls"] = urls
                    st.session_state["keyword_sources"] = keyword_sources
                else:
                    st.error("Aucun article trouvé. Veuillez vérifier les thèmes.")

            # Scraping
            scraped_data = []
            if urls:
                progress_bar = st.progress(0)
                status_text = st.empty()
                with st.spinner("Scraping des articles proposés..."):
                    for idx, url in enumerate(urls):
                        if len(scraped_data) >= 12:
                            break
                        page_data = scrape_page(url)
                        if page_data:
                            save_page_to_mongodb(page_data)
                            scraped_data.append(page_data)
                        else:
                            st.warning(f"Échec du scraping pour l'URL : {url}")
                        progress_bar.progress(min((idx + 1) / min(len(urls), 12), 1.0))
                        status_text.text(f"Scraping en cours... {min(idx + 1, 12)}/12")

                    if scraped_data:
                        st.session_state["scraped_data"] = scraped_data
                        st.success("Scraping des articles terminé avec succès.")
                    else:
                        st.error("Échec du scraping pour tous les articles.")

            # Génération des résumés
            if "scraped_data" in st.session_state:
                progress_bar = st.progress(0)
                status_text = st.empty()
                with st.spinner("Génération des résumés en cours..."):
                    summaries = []
                    for idx, (page, keyword) in enumerate(zip(st.session_state["scraped_data"], st.session_state["keyword_sources"])):
                        summary = generate_summary(
                            article_text=page.content,
                            system_prompt="""Role: Vous êtes un rédacteur expert en création de résumés d’articles clairs, informatifs et impartiaux.
                                            Votre objectif est de rédiger des résumés précis et concis qui permettent aux lecteurs de comprendre rapidement les points essentiels de l’article.
                                            Ces résumés doivent présenter les informations principales, les points clés et les conclusions importantes de manière neutre et fidèle au contenu de l’article, sans introduire d’éléments promotionnels, de suspense ou de langage intrigant.
                                            Adoptez un ton clair, adapté au sujet et accessible à un large public.
                                            Assurez-vous que le résumé est organisé de façon logique et structurée.""",
                            user_prompt=f"""Voici le contenu d’un article que je souhaite résumer :\n{page.content}
                                            Rédigez un résumé clair, concis et informatif de cet article. Mettez en avant les informations principales, les points clés et les conclusions importantes en vous assurant que :
                                            - Le ton est strictement neutre et descriptif.
                                            - La présentation est structurée et factuelle.
                                            - Le texte reste accessible et compréhensible.
                                            Ajoutez des émojis pertinents pour améliorer la lisibilité et rendre le résumé plus engageant, sans compromettre la neutralité du contenu. Limitez-vous à un maximum de 150 mots.""",
                        )
                        summaries.append(
                            {
                                "title": page.title,
                                "url": page.link,
                                "summary": summary,
                                "image_url": page.image_url,
                                "keyword_source": keyword,
                            }
                        )
                        progress_bar.progress((idx + 1) / len(st.session_state["scraped_data"]))
                        status_text.text(f"Génération des résumés en cours... {idx + 1}/{len(st.session_state['scraped_data'])}")

                    if summaries:
                        st.session_state["summaries"] = summaries
                        st.success("Génération des résumés terminée avec succès.")
                        save_results_to_file(input_data, "summaries", summaries)
                    else:
                        st.error("Aucun résumé n'a été généré.")

    if "summaries" in st.session_state:
        summaries = st.session_state["summaries"]
        st.write("### Articles proposés :")
        for i in range(0, len(summaries), 3):
            cols = st.columns(3)
            for j in range(3):
                if i + j < len(summaries):
                    summary = summaries[i + j]
                    with cols[j]:
                        link = summary.get('url', summary.get('link', '#'))
                        st.markdown(
                            f"<a href='{link}' style='text-decoration:none; color:inherit;'><h3>{summary['title']}</h3></a>",
                            unsafe_allow_html=True,
                        )
                        keyword_source = summary.get("keyword_source", "Source inconnue")
                        st.write(f"**Mot-clé source :** {keyword_source}")
                        if "image_url" in summary and summary["image_url"] and is_valid_image_url(summary["image_url"]):
                            st.markdown(
                                f"<a href='{link}'><img src='{summary['image_url']}' style='width:100%;'></a>",
                                unsafe_allow_html=True,
                            )
                        else:
                            st.warning("Image non disponible")
                        st.write(summary["summary"])
                        if link != '#':
                            st.markdown(
                                f"Pour en savoir plus, consultez l'article : [Article]({link})",
                                unsafe_allow_html=True,
                            )

                        # Boutons de feedback -> MySQL
                        if st.button("👍", key=f"like_{i + j}"):
                            save_feedback_to_mysql({
                                "Date": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                                "Onglet": "Suggestions d'articles par thèmes personnalisés",
                                "Unité de temps": time_unit,
                                "Titre réponse": summary['title'],
                                "Contenu réponse": summary["summary"],
                                "Réponse URL(s)": link,
                                "Avis utilisateur": "👍"
                            })
                            st.success("Votre avis a été enregistré !")

                        if st.button("👎", key=f"dislike_{i + j}"):
                            save_feedback_to_mysql({
                                "Date": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                                "Onglet": "Suggestions d'articles par thèmes personnalisés",
                                "Unité de temps": time_unit,
                                "Titre réponse": summary['title'],
                                "Contenu réponse": summary["summary"],
                                "Réponse URL(s)": link,
                                "Avis utilisateur": "👎"
                            })
                            st.success("Votre avis a été enregistré !")

    # Nouvelle recherche
    st.write("### Nouvelle recherche")
    new_keywords_input = st.text_area(
        "Entrez les nouveaux thèmes, séparés par une ligne",
        placeholder="thème1\nthème2",
    )
    new_keywords = [keyword.strip() for keyword in new_keywords_input.split("\n") if keyword.strip()]
    new_time_unit = st.selectbox("Unité de temps", ["mois", "années"], key="new_time_unit_tab1")
    new_time_value = st.number_input("Paru il y a", min_value=1, step=1, key="new_time_value_tab1")
    if st.button("Générer les résumés", key="generate_summaries_tab1"):
        new_input_data = "\n".join(new_keywords) + f"{new_time_unit}{new_time_value}"
        if check_and_load_results(new_input_data, "summaries"):
            st.success("Chargement des résultats précédents.")
            if "summaries" in st.session_state:
                new_summaries = st.session_state["summaries"]
                st.write("### Articles proposés pour la nouvelle recherche :")
                for i in range(0, len(new_summaries), 3):
                    cols = st.columns(3)
                    for j in range(3):
                        if i + j < len(new_summaries):
                            summary = new_summaries[i + j]
                            with cols[j]:
                                st.markdown(
                                    f"<a href='{summary['url']}' style='text-decoration:none; color:inherit;'><h3>{summary['title']}</h3></a>",
                                    unsafe_allow_html=True,
                                )
                                keyword_source = summary.get("keyword_source", "Source inconnue")
                                st.write(f"**Mot-clé source :** {keyword_source}")
                                if "image_url" in summary and summary["image_url"] and is_valid_image_url(summary["image_url"]):
                                    st.markdown(
                                        f"<a href='{summary['url']}'><img src='{summary['image_url']}' style='width:100%;'></a>",
                                        unsafe_allow_html=True,
                                    )
                                else:
                                    st.warning("Image non disponible")
                                st.write(summary["summary"])
                                if "url" in summary:
                                    st.markdown(
                                        f"Pour en savoir plus, consultez l'article : [Article]({summary['url']})",
                                        unsafe_allow_html=True,
                                    )
                                # Boutons de feedback -> MySQL
                                if st.button("👍", key=f"new_like_{i + j}"):
                                    save_feedback_to_mysql({
                                        "Date": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                                        "Onglet": "Suggestions d'articles par thèmes personnalisés",
                                        "Unité de temps": new_time_unit,
                                        "Titre réponse": summary['title'],
                                        "Contenu réponse": summary["summary"],
                                        "Réponse URL(s)": summary['url'],
                                        "Avis utilisateur": "👍"
                                    })
                                    st.success("Votre avis a été enregistré !")
                                if st.button("👎", key=f"new_dislike_{i + j}"):
                                    save_feedback_to_mysql({
                                        "Date": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                                        "Onglet": "Suggestions d'articles par thèmes personnalisés",
                                        "Unité de temps": new_time_unit,
                                        "Titre réponse": summary['title'],
                                        "Contenu réponse": summary["summary"],
                                        "Réponse URL(s)": summary['url'],
                                        "Avis utilisateur": "👎"
                                    })
                                    st.success("Votre avis a été enregistré !")
        else:
            if not new_keywords:
                st.error("Veuillez fournir au moins un mot-clé.")
            else:
                progress_bar = st.progress(0)
                status_text = st.empty()
                with st.spinner("Recherche des articles en cours..."):
                    new_urls = []
                    new_keyword_sources = []
                    for idx, keyword in enumerate(new_keywords):
                        keyword_urls = google_search(
                            query=keyword,
                            num_results=st.session_state.get("num_articles_keywords", 10),
                            time_unit=new_time_unit,
                            time_value=new_time_value,
                            exclude_ads=filters.get("exclude_ads", False),
                            exclude_professional=filters.get("exclude_professional", False),
                            target_press=filters.get("target_press", False),
                            exclude_jobs=filters.get("exclude_jobs", False),
                            exclude_training=filters.get("exclude_training", False),
                        )
                        new_urls.extend(keyword_urls)
                        new_keyword_sources.extend([keyword] * len(keyword_urls))
                        progress_bar.progress((idx + 1) / len(new_keywords))
                        status_text.text(f"Recherche en cours... {idx + 1}/{len(new_keywords)}")

                    if new_urls:
                        st.session_state["proposed_urls"] = new_urls
                        st.session_state["keyword_sources"] = new_keyword_sources
                    else:
                        st.error("Aucun article trouvé. Veuillez vérifier les thèmes.")

                # Scraping
                new_scraped_data = []
                if new_urls:
                    progress_bar = st.progress(0)
                    status_text = st.empty()
                    with st.spinner("Scraping des articles proposés..."):
                        for idx, url in enumerate(new_urls):
                            if len(new_scraped_data) >= 12:
                                break
                            page_data = scrape_page(url)
                            if page_data:
                                save_page_to_mongodb(page_data)
                                new_scraped_data.append(page_data)
                            else:
                                st.warning(f"Échec du scraping pour l'URL : {url}")
                            progress_bar.progress((idx + 1) / min(len(new_urls), 12))
                            status_text.text(f"Scraping en cours... {min(idx + 1, 12)}/12")

                        if new_scraped_data:
                            st.session_state["scraped_data"] = new_scraped_data
                            st.success("Scraping des articles terminé avec succès.")
                        else:
                            st.error("Échec du scraping pour tous les articles.")

                # Génération des résumés
                if "scraped_data" in st.session_state:
                    progress_bar = st.progress(0)
                    status_text = st.empty()
                    with st.spinner("Génération des résumés en cours..."):
                        new_summaries = []
                        for idx, (page, keyword) in enumerate(zip(st.session_state["scraped_data"], st.session_state["keyword_sources"])):
                            new_summary = generate_summary(
                                article_text=page.content,
                                system_prompt="""Role: Vous êtes un rédacteur expert en création de résumés d’articles clairs, informatifs et impartiaux.
                                                Votre objectif est de rédiger des résumés précis et concis qui permettent aux lecteurs de comprendre rapidement les points essentiels de l’article.
                                                Ces résumés doivent présenter les informations principales, les points clés et les conclusions importantes de manière neutre et fidèle au contenu de l’article, sans introduire d’éléments promotionnels, de suspense ou de langage intrigant.
                                                Adoptez un ton clair, adapté au sujet et accessible à un large public.
                                                Assurez-vous que le résumé est organisé de façon logique et structurée.""",
                                user_prompt=f"""Voici le contenu d’un article que je souhaite résumer :\n{page.content}
                                                Rédigez un résumé clair, concis et informatif de cet article. Mettez en avant les informations principales, les points clés et les conclusions importantes en vous assurant que :
                                                - Le ton est strictement neutre et descriptif.
                                                - La présentation est structurée et factuelle.
                                                - Le texte reste accessible et compréhensible.
                                                Ajoutez des émojis pertinents pour améliorer la lisibilité et rendre le résumé plus engageant, sans compromettre la neutralité du contenu. Limitez-vous à un maximum de 150 mots.""",
                            )
                            new_summaries.append(
                                {
                                    "title": page.title,
                                    "url": page.link,
                                    "summary": new_summary,
                                    "image_url": page.image_url,
                                    "keyword_source": keyword,
                                }
                            )
                            progress_bar.progress((idx + 1) / len(st.session_state["scraped_data"]))
                            status_text.text(f"Génération des résumés en cours... {idx + 1}/{len(st.session_state['scraped_data'])}")

                        if new_summaries:
                            st.session_state["summaries"] = new_summaries
                            st.success("Génération des résumés terminée avec succès.")
                            save_results_to_file(new_input_data, "summaries", new_summaries)
                        else:
                            st.error("Aucun résumé n'a été généré.")

##################################
# Onglet 3 : Suggestions d'articles par URLs sources
##################################
with tabs[2]:
    st.write("### Suggestions d'articles par URLs sources")
    source_urls = load_default_sources()
    filters = load_filters()
    time_unit = filters.get("time_unit", "mois")
    time_value = filters.get("time_value", 1)
    input_data = "\n".join(source_urls) + f"{time_unit}{time_value}"

    if check_and_load_results(input_data, "summaries"):
        st.success("Chargement des résultats précédents.")
    else:
        if not source_urls:
            st.error("Veuillez fournir au moins une URL source valide.")
        else:
            progress_bar = st.progress(0)
            status_text = st.empty()
            proposed_urls = []
            scraped_data = []
            with st.spinner("Scraping des pages sources en cours..."):
                for idx, url in enumerate(source_urls):
                    url = url.strip()
                    if not url or not url.startswith("http"):
                        st.warning(f"URL invalide ou vide : '{url}'")
                        continue
                    response = requests.get(url)
                    if response.status_code == 200:
                        soup = BeautifulSoup(response.content, "lxml")
                        articles = soup.select("h2.entry-title.ast-blog-single-element a, div.td-module-thumb a, div.tds_module_loop_1 a")
                        proposed_urls.extend([a["href"] for a in articles[:st.session_state.get("num_articles_sources", 10)]])
                    else:
                        st.error(f"Échec du scraping pour l'URL : {url}")
                    progress_bar.progress((idx + 1) / len(source_urls))
                    status_text.text(f"Scraping des pages sources en cours... {idx + 1}/{len(source_urls)}")

                if proposed_urls:
                    st.session_state["proposed_urls"] = proposed_urls
                    st.success("Scraping des pages sources terminé avec succès.")
                else:
                    st.error("Aucun article proposé trouvé. Veuillez vérifier les URLs sources.")

            # Scraping des articles proposés
            progress_bar = st.progress(0)
            status_text = st.empty()
            with st.spinner("Scraping des articles proposés en cours..."):
                for idx, url in enumerate(proposed_urls):
                    if len(scraped_data) >= 12:
                        break
                    page_data = scrape_page(url)
                    if page_data:
                        save_page_to_mongodb(page_data)
                        scraped_data.append(page_data)
                    else:
                        st.warning(f"Échec du scraping pour l'URL : {url}")
                    progress_bar.progress((idx + 1) / min(len(proposed_urls), 12))
                    status_text.text(f"Scraping des articles proposés en cours... {min(idx + 1, 12)}/12")

                if scraped_data:
                    st.session_state["scraped_data"] = scraped_data
                    st.success("Scraping des articles proposés terminé avec succès.")
                else:
                    st.error("Aucun article proposé n'a été scrapé. Veuillez vérifier les URLs.")

            # Génération des résumés
            progress_bar = st.progress(0)
            status_text = st.empty()
            with st.spinner("Génération des résumés en cours..."):
                summaries = []
                for idx, page in enumerate(scraped_data):
                    summary = generate_summary(
                        article_text=page.content,
                        system_prompt="""Role: Vous êtes un rédacteur expert en création de résumés d’articles clairs, informatifs et impartiaux.
                                        Votre objectif est de rédiger des résumés précis et concis qui permettent aux lecteurs de comprendre rapidement les points essentiels de l’article.
                                        Ces résumés doivent présenter les informations principales, les points clés et les conclusions importantes de manière neutre et fidèle au contenu de l’article, sans introduire d’éléments promotionnels, de suspense ou de langage intrigant.
                                        Adoptez un ton clair, adapté au sujet et accessible à un large public.
                                        Assurez-vous que le résumé est organisé de façon logique et structurée.""",
                        user_prompt=f"""Voici le contenu d’un article que je souhaite résumer :\n{page.content}
                                        Rédigez un résumé clair, concis et informatif de cet article. Mettez en avant les informations principales, les points clés et les conclusions importantes en vous assurant que :
                                        - Le ton est strictement neutre et descriptif.
                                        - La présentation est structurée et factuelle.
                                        - Le texte reste accessible et compréhensible.
                                        Ajoutez des émojis pertinents pour améliorer la lisibilité et rendre le résumé plus engageant, sans compromettre la neutralité du contenu. Limitez-vous à un maximum de 150 mots.""",
                    )
                    summaries.append(
                        {
                            "title": page.title,
                            "url": page.link,
                            "summary": summary,
                            "image_url": page.image_url,
                        }
                    )
                    progress_bar.progress((idx + 1) / len(scraped_data))
                    status_text.text(f"Génération des résumés en cours... {idx + 1}/{len(scraped_data)}")

                if summaries:
                    st.session_state["summaries"] = summaries
                    st.success("Génération des résumés terminée avec succès.")
                    save_results_to_file(input_data, "summaries", summaries)
                else:
                    st.error("Aucun résumé n'a été généré. Veuillez vérifier les articles.")

    if "summaries" in st.session_state:
        summaries = st.session_state["summaries"]
        st.write("### Articles proposés :")
        for i in range(0, len(summaries), 3):
            cols = st.columns(3)
            for j in range(3):
                if i + j < len(summaries):
                    summary = summaries[i + j]
                    with cols[j]:
                        link = summary.get('url', summary.get('link', '#'))
                        st.markdown(
                            f"<a href='{link}' style='text-decoration:none; color:inherit;'><h3>{summary['title']}</h3></a>",
                            unsafe_allow_html=True,
                        )
                        if "image_url" in summary and summary["image_url"] and is_valid_image_url(summary["image_url"]):
                            st.markdown(
                                f"<a href='{link}'><img src='{summary['image_url']}' style='width:100%;'></a>",
                                unsafe_allow_html=True,
                            )
                        else:
                            st.warning("Image non disponible")
                        st.write(summary["summary"])
                        if link != '#':
                            st.markdown(
                                f"Pour en savoir plus, consultez l'article : [Article]({link})",
                                unsafe_allow_html=True,
                            )
                        # Feedback -> MySQL
                        if st.button("👍", key=f"like_source_{i + j}"):
                            save_feedback_to_mysql({
                                "Date": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                                "Onglet": "Suggestions d'articles par URLs sources",
                                "Unité de temps": time_unit,
                                "Titre réponse": summary['title'],
                                "Contenu réponse": summary["summary"],
                                "Réponse URL(s)": link,
                                "Avis utilisateur": "👍"
                            })
                            st.success("Votre avis a été enregistré !")
                        if st.button("👎", key=f"dislike_source_{i + j}"):
                            save_feedback_to_mysql({
                                "Date": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                                "Onglet": "Suggestions d'articles par URLs sources",
                                "Unité de temps": time_unit,
                                "Titre réponse": summary['title'],
                                "Contenu réponse": summary["summary"],
                                "Réponse URL(s)": link,
                                "Avis utilisateur": "👎"
                            })
                            st.success("Votre avis a été enregistré !")

##################################
# Onglet 4 : Résumé automatique d'article
##################################
with tabs[3]:
    st.write("### Résumé(s) automatique(s) d'article(s)")
    st.write("Entrez une ou plusieurs URLs d'articles (une URL par ligne) pour générer des résumés.")
    urls_input = st.text_area(
        "Entrez le(s) URL(s) de(s) article(s) à résumer",
        placeholder="https://example1.com\nhttps://example2.com",
        value="\n".join(default_urls_summary),
        key="urls_input_summary_tab",
    )
    uploaded_files = st.file_uploader("Ou uploader des fichiers PDF ou Word", type=["pdf", "docx"], accept_multiple_files=True, key="file_uploader_summary_tab")
    urls = [url.strip() for url in urls_input.split("\n") if url.strip()]

    if st.button("Générer le(s) résumé(s)", key="generate_summaries_button"):
        input_data = "\n".join(urls)
        if check_and_load_results(input_data, "summaries"):
            st.success("Chargement des résultats précédents.")
        else:
            if not urls and not uploaded_files:
                st.error("Veuillez entrer au moins une URL valide ou uploader un fichier.")
            else:
                progress_bar = st.progress(0)
                status_text = st.empty()
                summaries = []

                # Scraping des URLs
                for idx, url in enumerate(urls):
                    try:
                        page_data = scrape_page(url)
                        if not page_data:
                            st.warning(f"Échec du scraping pour l'URL : {url}")
                            continue
                        save_page_to_mongodb(page_data)

                        system_prompt = """Vous êtes un expert en résumés d'articles destinés à un professionnel de l'innovation.
                        Votre tâche est de créer des résumés concis mais percutants pour différents types d'articles.
                        Chaque résumé doit mettre en avant les informations clés pertinentes en fonction du type d'article.
                        Les résumés sont destinés à un professionnel de l’innovation qui souhaite rester informé sur les solutions existantes et les tendances du marché."""
                        user_prompt = f"Résumez l'article suivant en respectant les consignes données :\n {page_data.content}"
                        st.session_state["system_prompt"] = system_prompt
                        st.session_state["user_prompt"] = user_prompt

                        with st.spinner(f"Génération du résumé pour {url}..."):
                            summary = generate_summary(
                                article_text=page_data.content,
                                system_prompt=system_prompt,
                                user_prompt=user_prompt,
                            )

                        if not summary or "Erreur" in summary:
                            st.error(f"Erreur lors de la génération du résumé pour l'URL : {url}")
                            continue

                        summaries.append(
                            {
                                "title": page_data.title,
                                "link": url,
                                "summary": summary,
                                "content": page_data.content,
                            }
                        )
                        st.success(f"Résumé généré pour : {page_data.title}")
                        progress_bar.progress((idx + 1) / len(urls))
                        status_text.text(f"Génération des résumés en cours... {idx + 1}/{len(urls)}")

                    except Exception as e:
                        st.error(f"Erreur inattendue pour l'URL {url} : {str(e)}")

                # Scraping des fichiers uploadés
                for idx, uploaded_file in enumerate(uploaded_files):
                    try:
                        if uploaded_file.name.endswith(".pdf"):
                            reader = PdfReader(uploaded_file)
                            content = "\n".join([page.extract_text() for page in reader.pages])
                        elif uploaded_file.name.endswith(".docx"):
                            doc = Document(uploaded_file)
                            content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                        else:
                            st.error(f"Type de fichier non supporté : {uploaded_file.name}")
                            continue

                        system_prompt = """Vous êtes un expert en résumés d'articles destinés à un professionnel de l'innovation.
                        Votre tâche est de créer des résumés concis mais percutants pour différents types d'articles."""
                        user_prompt = f"Résumez l'article suivant en respectant les consignes données :\n {content}"
                        st.session_state["system_prompt"] = system_prompt
                        st.session_state["user_prompt"] = user_prompt
                        with st.spinner(f"Génération du résumé pour {uploaded_file.name}..."):
                            summary = generate_summary(
                                article_text=content,
                                system_prompt=system_prompt,
                                user_prompt=user_prompt,
                            )
                        if not summary or "Erreur" in summary:
                            st.error(f"Erreur lors de la génération du résumé pour le fichier : {uploaded_file.name}")
                            continue

                        summaries.append(
                            {
                                "title": uploaded_file.name,
                                "link": "#",
                                "summary": summary,
                                "content": content,
                            }
                        )
                        st.success(f"Résumé généré pour : {uploaded_file.name}")
                        progress_bar.progress((idx + len(urls) + 1) / (len(urls) + len(uploaded_files)))
                        status_text.text(f"Génération des résumés en cours... {idx + len(urls) + 1}/{len(urls) + len(uploaded_files)}")

                    except Exception as e:
                        st.error(f"Erreur inattendue pour le fichier {uploaded_file.name} : {str(e)}")

                if summaries:
                    st.session_state["summaries"] = summaries
                    st.success("Génération des résumés terminée avec succès.")
                    save_results_to_file(input_data, "summaries", summaries)
                else:
                    st.warning("Aucun résumé n'a pu être généré. Veuillez vérifier les URLs fournies.")

    if "summaries" in st.session_state:
        summaries = st.session_state["summaries"]
        st.write("### Résumés générés :")
        for summary in summaries:
            link = summary.get('url', summary.get('link', '#'))
            st.markdown(
                f"<a href='{link}' style='text-decoration:none; color:inherit;'><h3>{summary['title']}</h3></a>",
                unsafe_allow_html=True,
            )
            st.write(summary["summary"])
            # Feedback -> MySQL
            if st.button("👍", key=f"like_summary_{summary['title']}"):
                save_feedback_to_mysql({
                    "Date": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                    "Onglet": "Résumé(s) automatique(s) d'article(s)",
                    "Unité de temps": "",
                    "Titre réponse": summary['title'],
                    "Contenu réponse": summary["summary"],
                    "Réponse URL(s)": link,
                    "Avis utilisateur": "👍"
                })
                st.success("Votre avis a été enregistré !")
            if st.button("👎", key=f"dislike_summary_{summary['title']}"):
                save_feedback_to_mysql({
                    "Date": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                    "Onglet": "Résumé(s) automatique(s) d'article(s)",
                    "Unité de temps": "",
                    "Titre réponse": summary['title'],
                    "Contenu réponse": summary["summary"],
                    "Réponse URL(s)": link,
                    "Avis utilisateur": "👎"
                })
                st.success("Votre avis a été enregistré !")

##################################
# Onglet 5 : Synthèse de corpus d'articles
##################################
with tabs[4]:
    st.write("### Synthèse de corpus d'articles")
    st.write("Fournissez les URLs de plusieurs articles pour générer une synthèse globale.")
    urls_input = st.text_area(
        "Entrez les URLs des articles, une par ligne",
        placeholder="https://example1.com\nhttps://example2.com",
        value="\n".join(default_urls_corpus),
    )
    uploaded_files = st.file_uploader("Ou uploader des fichiers PDF ou Word", type=["pdf", "docx"], accept_multiple_files=True, key="file_uploader_corpus_tab")

    urls = [url.strip() for url in urls_input.split("\n") if url.strip()]
    if st.button("Générer la synthèse de corpus"):
        input_data = "\n".join(urls)
        if check_and_load_results(input_data, "synthesis"):
            st.success("Chargement des résultats précédents.")
        elif not urls and not uploaded_files:
            st.error("Veuillez fournir au moins une URL valide ou uploader un fichier.")
        else:
            progress_bar = st.progress(0)
            status_text = st.empty()
            scraped_data = []

            # Scraping des URLs
            for idx, url in enumerate(urls):
                page_data = scrape_page(url)
                if page_data:
                    save_page_to_mongodb(page_data)
                    scraped_data.append(page_data)
                else:
                    st.warning(f"Échec du scraping pour l'URL : {url}")
                progress_bar.progress((idx + 1) / len(urls))
                status_text.text(f"Scraping des articles en cours... {idx + 1}/{len(urls)}")

            # Scraping des fichiers
            for uploaded_file in uploaded_files:
                try:
                    if uploaded_file.name.endswith(".pdf"):
                        reader = PdfReader(uploaded_file)
                        content = "\n".join([page.extract_text() for page in reader.pages])
                    elif uploaded_file.name.endswith(".docx"):
                        doc = Document(uploaded_file)
                        content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                    else:
                        st.error(f"Type de fichier non supporté : {uploaded_file.name}")
                        continue

                    page = Page(
                        date=None,
                        title=uploaded_file.name,
                        link="#",
                        description="",
                        content=content,
                        author=None,
                        image_url=None,
                    )
                    scraped_data.append(page)
                    progress_bar.progress((idx + len(urls) + 1) / (len(urls) + len(uploaded_files)))
                    status_text.text(f"Scraping des articles en cours... {idx + len(urls) + 1}/{len(urls) + len(uploaded_files)}")
                except Exception as e:
                    st.error(f"Erreur inattendue pour le fichier {uploaded_file.name} : {str(e)}")

            if not scraped_data:
                st.error("Aucun article n'a été scrappé. Veuillez vérifier les URLs fournies.")
            else:
                concatenated_content = "\n\n".join(
                    f"### {page.title}\nURL : {page.link}\n{page.content}"
                    for page in scraped_data
                )
                with st.spinner("Génération de la synthèse..."):
                    try:
                        system_prompt = """Vous êtes un expert en veille stratégique. Votre tâche est de créer une synthèse concise et engageante des articles suivants. 
                        1. Commencez par une vue d'ensemble des grandes tendances observées.
                        2. Pour chaque article, identifiez les points clés et les innovations notables.
                        3. Concluez avec une analyse des implications globales."""
                        user_prompt = f"""Veuillez générer une synthèse stratégique des articles suivants en suivant ces consignes :
                        {concatenated_content}"""
                        synthesis = generate_summary(
                            article_text=concatenated_content,
                            system_prompt=system_prompt,
                            user_prompt=user_prompt,
                        )

                        if "Erreur" in synthesis:
                            st.error(synthesis)
                        else:
                            st.success("Synthèse générée avec succès !")
                            st.write("### Synthèse du corpus")
                            st.write(synthesis)
                            st.session_state["synthesis"] = synthesis
                            file = create_file(
                                summary=synthesis,
                                article_text=concatenated_content,
                                system_prompt=system_prompt,
                                user_prompt=user_prompt,
                                url="Synthèse du corpus",
                                title="Synthèse de plusieurs articles",
                            )
                            st.download_button(
                                label="Télécharger la synthèse",
                                data=file,
                                file_name="synthese_corpus.pdf",
                                mime="application/pdf",
                            )
                            save_results_to_file(input_data, "synthesis", {"synthesis": synthesis})

                            # Feedback -> MySQL
                            if st.button("👍", key=f"like_synthesis"):
                                save_feedback_to_mysql({
                                    "Date": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                                    "Onglet": "Synthèse de corpus d'articles",
                                    "Unité de temps": "",
                                    "Titre réponse": "Synthèse de plusieurs articles",
                                    "Contenu réponse": synthesis,
                                    "Réponse URL(s)": "Synthèse du corpus",
                                    "Avis utilisateur": "👍"
                                })
                                st.success("Votre avis a été enregistré !")

                            if st.button("👎", key=f"dislike_synthesis"):
                                save_feedback_to_mysql({
                                    "Date": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                                    "Onglet": "Synthèse de corpus d'articles",
                                    "Unité de temps": "",
                                    "Titre réponse": "Synthèse de plusieurs articles",
                                    "Contenu réponse": synthesis,
                                    "Réponse URL(s)": "Synthèse du corpus",
                                    "Avis utilisateur": "👎"
                                })
                                st.success("Votre avis a été enregistré !")

                    except Exception as e:
                        st.error(f"Erreur inattendue : {str(e)}")

    if "synthesis" in st.session_state:
        st.write("### Dernière synthèse du corpus")
        synthesis_data = st.session_state["synthesis"]
        if isinstance(synthesis_data, dict):
            synthesis_text = synthesis_data.get("synthesis", "")
        else:
            synthesis_text = synthesis_data

        if synthesis_text:
            st.markdown(synthesis_text)
        else:
            st.error("La synthèse est vide ou non disponible.")

##################################
# Onglet 6 : Chatbot Q/A
##################################
with tabs[5]:
    st.write("### Chabot Q/A")
    st.write("Chargez des articles via des URLs, des fichiers PDF ou Word, et posez vos questions.")

    # Chargement des articles
    urls_input = st.text_area(
        "Entrez les URLs des articles, une par ligne",
        placeholder="https://example1.com\nhttps://example2.com",
        key="urls_input_qa_tab_unique",
    )
    uploaded_files = st.file_uploader("Ou uploader des fichiers PDF ou Word", type=["pdf", "docx"], accept_multiple_files=True, key="file_uploader_qa_tab_unique")

    urls = [url.strip() for url in urls_input.split("\n") if url.strip()]

    if st.button("Charger les articles", key="load_articles_button_unique"):
        scraped_data = []
        if urls:
            progress_bar = st.progress(0)
            status_text = st.empty()
            with st.spinner("Chargement des articles en cours..."):
                for idx, url in enumerate(urls):
                    page_data = scrape_page(url)
                    if page_data:
                        save_page_to_mongodb(page_data)
                        scraped_data.append(page_data)
                    else:
                        st.warning(f"Échec du scraping pour l'URL : {url}")
                    progress_bar.progress((idx + 1) / len(urls))
                    status_text.text(f"Chargement des articles en cours... {idx + 1}/{len(urls)}")

        if uploaded_files:
            progress_bar = st.progress(0)
            status_text = st.empty()
            with st.spinner("Chargement des fichiers en cours..."):
                for idx, uploaded_file in enumerate(uploaded_files):
                    try:
                        if uploaded_file.name.endswith(".pdf"):
                            reader = PdfReader(uploaded_file)
                            content = "\n".join([page.extract_text() for page in reader.pages])
                        elif uploaded_file.name.endswith(".docx"):
                            doc = Document(uploaded_file)
                            content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                        else:
                            st.error(f"Type de fichier non supporté : {uploaded_file.name}")
                            continue

                        page = Page(
                            date=None,
                            title=uploaded_file.name,
                            link="#",
                            description="",
                            content=content,
                            author=None,
                            image_url=None,
                        )
                        scraped_data.append(page)
                        progress_bar.progress((idx + 1) / len(uploaded_files))
                        status_text.text(f"Chargement des fichiers en cours... {idx + 1}/{len(uploaded_files)}")

                    except Exception as e:
                        st.error(f"Erreur inattendue pour le fichier {uploaded_file.name} : {str(e)}")

        if scraped_data:
            st.session_state["scraped_data"] = scraped_data
            st.success("Articles chargés avec succès.")
        else:
            st.error("Aucun article n'a été chargé. Veuillez vérifier les URLs et les fichiers fournis.")

    # Interface conversationnelle
    if "scraped_data" in st.session_state:
        st.write("### Posez vos questions")
        user_input = st.text_input("Votre question :", key="user_input_qa_tab_unique")
        if st.button("Envoyer", key="send_button_qa_tab_unique"):
            if user_input:
                concatenated_content = "\n\n".join(
                    f"### {page.title}\nURL : {page.link}\n{page.content}"
                    for page in st.session_state["scraped_data"]
                )
                with st.spinner("Génération de la réponse en cours..."):
                    answer = generate_answer(question=user_input, context=concatenated_content)

                if "conversation_history" not in st.session_state:
                    st.session_state["conversation_history"] = []
                st.session_state["conversation_history"].append({"user": user_input, "bot": answer})
            else:
                st.warning("Veuillez entrer une question.")
    else:
        st.warning("Veuillez d'abord charger des articles.")



    # Affichage de l'historique des conversations
    if "conversation_history" in st.session_state:
        for chat in st.session_state["conversation_history"]:
            with st.container():
                st.markdown(f'<div class="user-message">{chat["user"]}</div>', unsafe_allow_html=True)
            with st.container():
                st.markdown(f'<div class="bot-message">{chat["bot"]}</div>', unsafe_allow_html=True)
            st.write("---")
            # Feedback -> MySQL
            if st.button("👍", key=f"like_chat_{chat['user']}"):
                save_feedback_to_mysql({
                    "Date": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                    "Onglet": "Chatbot Q/A",
                    "Unité de temps": "",
                    "Titre réponse": chat['user'],
                    "Contenu réponse": chat["bot"],
                    "Réponse URL(s)": "",
                    "Avis utilisateur": "👍"
                })
                st.success("Votre avis a été enregistré !")
            if st.button("👎", key=f"dislike_chat_{chat['user']}"):
                save_feedback_to_mysql({
                    "Date": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                    "Onglet": "Chatbot Q/A",
                    "Unité de temps": "",
                    "Titre réponse": chat['user'],
                    "Contenu réponse": chat["bot"],
                    "Réponse URL(s)": "",
                    "Avis utilisateur": "👎"
                })
                st.success("Votre avis a été enregistré !")

    # Indicateur de chargement
    if st.session_state.get("loading", False):
        st.markdown('<div class="loading">...</div>', unsafe_allow_html=True)
